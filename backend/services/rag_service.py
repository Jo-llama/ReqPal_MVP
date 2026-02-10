import hashlib
import json
import csv
from pathlib import Path
from typing import List, Dict, Any
from io import BytesIO
import re

import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx

from models import (
    Document,
    DocumentChunk,
    DocumentClassification,
    RAGQuery,
    RAGResult,
)


class RAGService:
    """
    RAG Service with ChromaDB (persistent)

    IMPORTANT:
    - semantic_search() returns top_k candidates ALWAYS (no min_similarity filtering)
    - thresholding is done at answer layer
    """

    HEADING_RE = re.compile(r"^(\d+(\.\d+)*)\s+.+|^(SECTION|CHAPTER)\s+\w+|^[A-Z][A-Z\s\-]{6,}$")

    def __init__(self):
        print("📥 Loading embedding model...")
        # strong small english embeddings
        self.embedding_model = SentenceTransformer("BAAI/bge-small-en-v1.5")
        print("✅ Embedding model loaded")

        chroma_path = Path("storage/chroma")
        chroma_path.mkdir(parents=True, exist_ok=True)

        self.chroma_client = chromadb.PersistentClient(
            path=str(chroma_path),
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True,
            ),
        )

        # cosine distance
        self.collection = self.chroma_client.get_or_create_collection(
            name="reqpal_documents",
            metadata={"hnsw:space": "cosine"},
        )

        print("📦 Chroma collection metadata:", self.collection.metadata)

        self.upload_dir = Path("uploads")
        self.upload_dir.mkdir(exist_ok=True)

    # =========================================================
    # Document processing
    # =========================================================

    async def process_document(
            self,
            file_content: bytes,
            filename: str,
            project_id: int,
            classification: DocumentClassification,
            document_id: int,  # ✅ NEW
            metadata: Dict[str, Any] = {},
    ) -> Document:
        file_ext = Path(filename).suffix.lower().replace(".", "")
        file_hash = hashlib.md5(file_content).hexdigest()
        safe_name = f"{project_id}_{file_hash}_{filename}"
        file_path = self.upload_dir / safe_name

        with open(file_path, "wb") as f:
            f.write(file_content)

        # ✅ set ID immediately
        document = Document(
            id=document_id,  # ✅ NEW
            project_id=project_id,
            filename=filename,
            file_type=file_ext,
            file_path=str(file_path),
            file_size=len(file_content),
            classification=classification,
            processing_status="processing",
            metadata=metadata,
        )

        text = await self._extract_text(file_content, file_ext)
        chunks = await self._chunk_text(text, document.id, filename)  # ✅ not 0

        await self._embed_and_store(chunks, project_id, document)

        document.chunks = chunks
        document.total_chunks = len(chunks)
        document.indexed = True
        document.processing_status = "completed"

        return document

    async def _extract_text(self, content: bytes, file_type: str) -> str:
        if file_type == "pdf":
            reader = PyPDF2.PdfReader(BytesIO(content))
            pages = []
            for i, page in enumerate(reader.pages):
                t = page.extract_text() or ""
                if t.strip():
                    pages.append(f"[Page {i+1}]\n{t}")
            return "\n\n".join(pages)

        if file_type == "docx":
            d = docx.Document(BytesIO(content))
            return "\n\n".join(p.text for p in d.paragraphs if (p.text or "").strip())

        if file_type == "txt":
            return content.decode("utf-8", errors="ignore")

        if file_type == "csv":
            txt = content.decode("utf-8", errors="ignore")
            return "\n".join(" | ".join(r) for r in csv.reader(txt.splitlines()))

        if file_type == "json":
            return json.dumps(json.loads(content.decode("utf-8")), indent=2)

        raise ValueError(f"Unsupported file type: {file_type}")

    def _is_heading(self, line: str) -> bool:
        line = (line or "").strip()
        return bool(self.HEADING_RE.match(line))

    async def _chunk_text(
        self,
        text: str,
        document_id: int,
        filename: str,
        chunk_size: int = 1300,   # chars
        overlap: int = 240,       # chars
    ) -> List[DocumentChunk]:
        """
        Heading-aware chunking:
        1) Build blocks by headings + paragraph-ish boundaries
        2) Pack blocks into size-limited chunks with char-overlap
        """

        lines = [l.rstrip() for l in text.splitlines()]
        blocks: List[tuple[str | None, str]] = []
        buf: List[str] = []
        current_heading: str | None = None

        def flush_buf():
            nonlocal buf
            if buf:
                blocks.append((current_heading, "\n".join(buf).strip()))
                buf = []

        # 1) build blocks
        for line in lines:
            if not (line or "").strip():
                continue

            if self._is_heading(line):
                flush_buf()
                current_heading = line.strip()
                continue

            buf.append(line)

            # paragraph-ish boundary
            if line.endswith(".") and len(" ".join(buf)) > 280:
                flush_buf()

        flush_buf()

        # 2) pack blocks
        chunks: List[DocumentChunk] = []
        chunk_index = 0
        cur = ""
        cur_heading: str | None = None

        def add_chunk(content: str, heading: str | None):
            nonlocal chunk_index
            chunk_id = f"doc{document_id}_chunk{chunk_index}"
            chunks.append(
                DocumentChunk(
                    chunk_id=chunk_id,
                    document_id=document_id,
                    content=content.strip(),
                    chunk_index=chunk_index,
                    section=heading,
                    metadata={"filename": filename, "total_length": len(content)},
                )
            )
            chunk_index += 1

        for heading, block in blocks:
            block_text = f"{heading}\n{block}" if heading else block

            if len(cur) + len(block_text) + 2 > chunk_size and cur.strip():
                add_chunk(cur, cur_heading)

                tail = cur[-overlap:] if len(cur) > overlap else cur
                cur = (tail + "\n\n" + block_text).strip()
                cur_heading = heading or cur_heading
            else:
                if not cur_heading and heading:
                    cur_heading = heading
                cur = (cur + "\n\n" + block_text).strip() if cur else block_text

        if cur.strip():
            add_chunk(cur, cur_heading)

        return chunks

    async def _embed_and_store(
        self,
        chunks: List[DocumentChunk],
        project_id: int,
        document: Document,
    ):
        if not chunks:
            return

        texts = [c.content for c in chunks]
        ids = [c.chunk_id for c in chunks]

        # ✅ normalize for cosine (critical)
        embeddings = self.embedding_model.encode(
            texts,
            show_progress_bar=False,
            normalize_embeddings=True,
        )

        metadatas = [
            {
                "project_id": int(project_id),
                "document_id": int(document.id),
                "document_name": document.filename,
                "classification": document.classification.value,
                "chunk_index": int(c.chunk_index),
                "filename": c.metadata.get("filename"),
            }
            for c in chunks
        ]

        self.collection.add(
            ids=ids,
            documents=texts,
            embeddings=embeddings.tolist(),
            metadatas=metadatas,
        )

        for c in chunks:
            c.embedding_id = c.chunk_id

        print(f"💾 Indexed {len(chunks)} chunks")

    # =========================================================
    # Semantic search (no thresholding here)
    # =========================================================

    async def semantic_search(self, query: RAGQuery) -> List[RAGResult]:
        print(f"🔍 RAG search: '{query.query}'")

        q_emb = self.embedding_model.encode(
            [query.query],
            normalize_embeddings=True,
        )[0]

        where_filter: Dict[str, Any] = {"project_id": query.project_id}
        if query.document_classifications:
            where_filter["classification"] = {"$in": [c.value for c in query.document_classifications]}

        results = self.collection.query(
            query_embeddings=[q_emb.tolist()],
            n_results=query.top_k,
            where=where_filter,
        )

        ids = results["ids"][0] if results.get("ids") else []
        metadatas = results["metadatas"][0] if results.get("metadatas") else []
        docs = results["documents"][0] if results.get("documents") else []
        distances = results["distances"][0] if results.get("distances") else []

        print(f"DEBUG where_filter = {where_filter}")
        print(f"DEBUG got ids = {len(ids)}")
        print(f"DEBUG distances[:10] = {distances[:10]}")

        out: List[RAGResult] = []

        # cosine distance -> cosine similarity
        for i, chunk_id in enumerate(ids):
            md = metadatas[i] or {}
            content = docs[i] or ""
            dist = float(distances[i]) if i < len(distances) else 0.0

            sim = 1.0 - dist  # [-1..1] possible

            out.append(
                RAGResult(
                    chunk_id=chunk_id,
                    document_id=int(md.get("document_id") or 0),
                    document_name=str(md.get("document_name", "Unknown")),
                    content=content,
                    similarity_score=float(sim),
                    classification=str(md.get("classification", "unknown")),
                    page_number=md.get("page_number"),
                    section=md.get("section") or md.get("section"),
                )
            )

        return out

    def get_stats(self) -> Dict[str, Any]:
        return {
            "vector_db": "ChromaDB",
            "collection": "reqpal_documents",
            "chunks": self.collection.count(),
            "embedding_model": "BAAI/bge-small-en-v1.5",
            "space": self.collection.metadata.get("hnsw:space"),
        }


rag_service = RAGService()
__all__ = ["rag_service", "RAGService"]
